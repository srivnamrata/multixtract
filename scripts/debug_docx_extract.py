from docx import Document
from docx.enum.text import WD_BREAK
from PIL import Image
from pathlib import Path
import zipfile
from multixtract.extractors.docx import DocxExtractor
from multixtract.filters import ImageFilterPipeline

p = Path('C:/Users/Z0019386/AppData/Local/Temp/test_docx_debug')
p.mkdir(parents=True, exist_ok=True)
img_path = p / 'img.png'
img = Image.new('RGB', (300,300), color=(255,0,0))
img.save(img_path, format='PNG')

doc = Document()
p1 = doc.add_paragraph('Before page')
run = p1.add_run()
run.add_break(WD_BREAK.PAGE)

table = doc.add_table(rows=2, cols=2)
table.cell(0,0).text = 'A'

doc.add_picture(str(img_path))
doc_path = p / 'test.docx'
doc.save(str(doc_path))

print('Created docx at', doc_path)
with zipfile.ZipFile(str(doc_path), 'r') as zf:
    print('ZIP names:', zf.namelist())

extractor = DocxExtractor()
permissive = ImageFilterPipeline(min_image_size=10, min_image_size_minor=5)
doc_obj, prepared = extractor.extract(str(doc_path), image_filter=permissive)
print('Document metadata page_count:', doc_obj.get('metadata',{}).get('page_count'))
print('Prepared images count:', len(prepared))
if prepared:
    print('First prepared keys:', prepared[0].keys())
else:
    print('No prepared images')
