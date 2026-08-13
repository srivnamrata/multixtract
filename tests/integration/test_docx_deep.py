from PIL import Image


def test_docx_extractor_integration_with_table_image_and_pagebreak(tmp_path):
    from docx import Document as DocxDocument
    from docx.enum.text import WD_BREAK

    # Create an image with variation so it passes the solid-color filter
    img_path = tmp_path / "img.png"
    img = Image.new("RGB", (300, 300), color=(255, 0, 0))
    pixels = img.load()
    for x in range(0, 300, 10):
        for y in range(0, 300, 10):
            pixels[x, y] = (x % 255, y % 255, (x + y) % 255)
    img.save(img_path, format="PNG")

    doc = DocxDocument()
    p = doc.add_paragraph("Before page")
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"

    doc.add_picture(str(img_path))

    doc_path = tmp_path / "test.docx"
    doc.save(str(doc_path))

    from multixtract.extractors.docx import DocxExtractor
    from multixtract.filters import ImageFilterPipeline

    extractor = DocxExtractor()
    # Use a permissive filter to avoid rejecting our test image
    permissive = ImageFilterPipeline(min_image_size=10, min_image_size_minor=5)
    document, prepared_images = extractor.extract(str(doc_path), image_filter=permissive)

    assert document["metadata"]["page_count"] >= 1
    assert any(p["tables"] for p in document["pgs"]) or document["metadata"]["table_count"] >= 1
    assert len(prepared_images) >= 1
